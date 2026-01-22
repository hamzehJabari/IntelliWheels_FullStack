"""
IntelliWheels Documentation Converter
=====================================
Converts all Markdown (.md) documentation files to Word (.docx) format.

Prerequisites:
    pip install python-docx markdown

Usage:
    python convert_docs_to_word.py

This script converts all .md files in the docs/ folder to .docx format.
"""

import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx not installed. Install with: pip install python-docx")

def markdown_to_docx(md_content, output_path, title="Document"):
    """Convert markdown content to a Word document."""
    if not DOCX_AVAILABLE:
        print("Cannot convert - python-docx not installed")
        return False
    
    doc = Document()
    
    # Add title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), 1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), 2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), 3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), 4)
        
        # Horizontal rule
        elif line.strip() == '---':
            doc.add_paragraph('─' * 50)
        
        # Code blocks
        elif line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                code_para = doc.add_paragraph()
                code_run = code_para.add_run('\n'.join(code_lines))
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(9)
        
        # Tables (simplified - just add as text)
        elif line.strip().startswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                if not lines[i].strip().startswith('|---'):
                    table_lines.append(lines[i])
                i += 1
            i -= 1  # Back up one since we'll increment at the end
            
            # Parse table
            if len(table_lines) > 0:
                # Get column count from header
                header_cells = [c.strip() for c in table_lines[0].split('|') if c.strip()]
                col_count = len(header_cells)
                
                if col_count > 0:
                    table = doc.add_table(rows=len(table_lines), cols=col_count)
                    table.style = 'Table Grid'
                    
                    for row_idx, table_line in enumerate(table_lines):
                        cells = [c.strip() for c in table_line.split('|') if c.strip()]
                        for col_idx, cell_content in enumerate(cells[:col_count]):
                            if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                                table.rows[row_idx].cells[col_idx].text = cell_content
        
        # Bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            bullet_text = line.strip()[2:]
            doc.add_paragraph(bullet_text, style='List Bullet')
        
        # Numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            list_text = re.sub(r'^\d+\.\s', '', line.strip())
            doc.add_paragraph(list_text, style='List Number')
        
        # Checkboxes
        elif line.strip().startswith('☐') or line.strip().startswith('☑') or line.strip().startswith('✅'):
            doc.add_paragraph(line.strip(), style='List Bullet')
        
        # Bold text handling in regular paragraphs
        else:
            para = doc.add_paragraph()
            # Simple bold handling
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                else:
                    para.add_run(part)
        
        i += 1
    
    # Save document
    doc.save(output_path)
    return True


def convert_all_docs():
    """Convert all markdown files in docs/ to Word format."""
    docs_dir = Path(__file__).parent
    
    # List of files to convert
    md_files = [
        'STEP_BY_STEP_ACTION_PLAN.md',
        'FUTURE_ROADMAP.md',
        'BUSINESS_PLAN.md',
        'BUSINESS_READINESS_CHECKLIST.md',
        'PITCH_DECK_OUTLINE.md',
        'FULL_DOCUMENTATION.md',
        'VISUAL_DIAGRAMS.md',
    ]
    
    # Also check legal folder
    legal_files = [
        'legal/TERMS_OF_SERVICE.md',
        'legal/PRIVACY_POLICY.md',
        'legal/DEALER_AGREEMENT.md',
    ]
    
    all_files = md_files + legal_files
    
    print("=" * 60)
    print("IntelliWheels Documentation Converter")
    print("=" * 60)
    
    if not DOCX_AVAILABLE:
        print("\n❌ python-docx is not installed!")
        print("   Run: pip install python-docx")
        print("\nAlternative: Use online converters or Pandoc:")
        print("   pandoc input.md -o output.docx")
        return
    
    converted = 0
    failed = 0
    
    for md_file in all_files:
        md_path = docs_dir / md_file
        
        if not md_path.exists():
            print(f"⚠️  Skipping {md_file} - file not found")
            continue
        
        # Create output path
        docx_file = md_file.replace('.md', '.docx')
        docx_path = docs_dir / docx_file
        
        # Ensure output directory exists
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        
        print(f"\n📄 Converting: {md_file}")
        
        try:
            with open(md_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # Get title from first line
            title = md_file.replace('.md', '').replace('_', ' ').replace('-', ' ')
            first_line = md_content.split('\n')[0]
            if first_line.startswith('# '):
                title = first_line[2:].strip()
            
            success = markdown_to_docx(md_content, str(docx_path), title)
            
            if success:
                print(f"   ✅ Created: {docx_file}")
                converted += 1
            else:
                print(f"   ❌ Failed: {docx_file}")
                failed += 1
                
        except Exception as e:
            print(f"   ❌ Error: {str(e)}")
            failed += 1
    
    print("\n" + "=" * 60)
    print(f"Conversion Complete!")
    print(f"   ✅ Converted: {converted} files")
    print(f"   ❌ Failed: {failed} files")
    print("=" * 60)


if __name__ == '__main__':
    convert_all_docs()
