"""
Convert PROJECT_EXPLANATION.md to a Word document (.docx)

Usage:
    pip install python-docx
    python convert_explanation_to_docx.py
"""

import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT


def parse_markdown(md_content):
    """Parse markdown content into structured elements."""
    lines = md_content.split('\n')
    elements = []
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            elements.append(('h1', line[2:].strip()))
            i += 1
            continue
        elif line.startswith('## '):
            elements.append(('h2', line[3:].strip()))
            i += 1
            continue
        elif line.startswith('### '):
            elements.append(('h3', line[4:].strip()))
            i += 1
            continue
        elif line.startswith('#### '):
            elements.append(('h4', line[5:].strip()))
            i += 1
            continue
        
        # Horizontal rule
        if line.strip() == '---':
            elements.append(('hr', ''))
            i += 1
            continue
        
        # Code blocks
        if line.strip().startswith('```'):
            code_lines = []
            lang = line.strip()[3:]
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            elements.append(('code', '\n'.join(code_lines)))
            i += 1
            continue
        
        # Tables
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            elements.append(('table', table_lines))
            continue
        
        # Bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            bullet_text = line.strip()[2:]
            elements.append(('bullet', bullet_text))
            i += 1
            continue
        
        # Numbered lists
        if re.match(r'^\d+\.\s', line.strip()):
            list_text = re.sub(r'^\d+\.\s', '', line.strip())
            elements.append(('number', list_text))
            i += 1
            continue
        
        # Bold Q&A format
        if line.strip().startswith('**Q:'):
            elements.append(('question', line.strip()))
            i += 1
            continue
        if line.strip().startswith('> A:'):
            elements.append(('answer', line.strip()[2:]))
            i += 1
            continue
        
        # Regular paragraph
        elements.append(('para', line.strip()))
        i += 1
    
    return elements


def clean_markdown_formatting(text):
    """Remove markdown formatting and return plain text with formatting info."""
    # Remove inline code backticks
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Remove bold
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    # Remove italic
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    # Remove links but keep text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    return text


def add_formatted_text(paragraph, text):
    """Add text with basic formatting to a paragraph."""
    # Split by bold markers
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif '`' in part:
            # Handle inline code
            code_parts = re.split(r'(`[^`]+`)', part)
            for cp in code_parts:
                if cp.startswith('`') and cp.endswith('`'):
                    run = paragraph.add_run(cp[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                else:
                    paragraph.add_run(cp)
        else:
            paragraph.add_run(part)


def create_docx(md_path, output_path):
    """Convert markdown to Word document."""
    
    # Read markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Parse
    elements = parse_markdown(md_content)
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Process elements
    for elem_type, content in elements:
        
        if elem_type == 'h1':
            p = doc.add_heading(clean_markdown_formatting(content), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif elem_type == 'h2':
            doc.add_heading(clean_markdown_formatting(content), level=1)
            
        elif elem_type == 'h3':
            doc.add_heading(clean_markdown_formatting(content), level=2)
            
        elif elem_type == 'h4':
            doc.add_heading(clean_markdown_formatting(content), level=3)
            
        elif elem_type == 'hr':
            p = doc.add_paragraph()
            p.add_run('─' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif elem_type == 'code':
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.5)
            
        elif elem_type == 'table':
            # Parse table
            rows_data = []
            for line in content:
                if '---' in line:
                    continue
                cells = [c.strip() for c in line.split('|')[1:-1]]
                if cells:
                    rows_data.append(cells)
            
            if rows_data:
                table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
                table.style = 'Table Grid'
                
                for i, row_data in enumerate(rows_data):
                    for j, cell_text in enumerate(row_data):
                        cell = table.rows[i].cells[j]
                        cell.text = clean_markdown_formatting(cell_text)
                        # Bold header row
                        if i == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
                
                doc.add_paragraph()  # Space after table
                
        elif elem_type == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, content)
            
        elif elem_type == 'number':
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, content)
            
        elif elem_type == 'question':
            p = doc.add_paragraph()
            # Extract question text
            q_text = content.replace('**Q:', 'Q:').replace('**', '')
            run = p.add_run(q_text)
            run.bold = True
            
        elif elem_type == 'answer':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            add_formatted_text(p, content)
            
        elif elem_type == 'para':
            if content:
                p = doc.add_paragraph()
                add_formatted_text(p, content)
    
    # Save
    doc.save(output_path)
    print(f"✅ Created: {output_path}")


if __name__ == '__main__':
    docs_dir = Path(__file__).parent
    md_file = docs_dir / 'PROJECT_EXPLANATION.md'
    docx_file = docs_dir / 'PROJECT_EXPLANATION.docx'
    
    if not md_file.exists():
        print(f"❌ File not found: {md_file}")
        exit(1)
    
    create_docx(md_file, docx_file)
    print(f"\n📄 Word document saved to: {docx_file}")
